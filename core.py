# -*- coding: utf-8 -*-
import codecs
import logging
import os
import re
from os.path import abspath
try:
    from os import scandir
except ImportError:
    from scandir import scandir

logger = logging.getLogger(__name__)

DEFAULT_INDIRS = ['.']
DEFAULT_EXTS = ['py']
DEFAULT_COMMENT_CHARS = (
    '#', '//'
)
DEFAULT_SKIP_DIRS = [
    'node_modules', 'uni_modules', '__pycache__', '.git', '.venv', 'venv'
]
DEFAULT_SKIP_FILES = [
    'package.json', 'package-lock.json', 'pnpm-lock.yaml', 'yarn.lock'
]
LANGUAGE_BY_EXT = {
    'py': 'python',
    'js': 'javascript',
    'jsx': 'javascript',
    'ts': 'typescript',
    'tsx': 'typescript',
    'go': 'go',
    'php': 'php',
    'cs': 'csharp',
    'kt': 'kotlin',
    'kts': 'kotlin',
    'swift': 'swift',
    'rs': 'rust',
    'dart': 'dart',
    'scala': 'scala',
    'sql': 'sql',
    'r': 'r',
    'lua': 'lua',
    'ps1': 'powershell',
    'psm1': 'powershell',
    'json': 'json',
    'yaml': 'yaml',
    'yml': 'yaml',
    'java': 'java',
    'c': 'c',
    'h': 'c',
    'cpp': 'cpp',
    'hpp': 'cpp',
    'cc': 'cpp',
    'html': 'html',
    'htm': 'html',
    'xml': 'xml',
    'css': 'css',
    'sh': 'shellscript',
    'bash': 'shellscript',
    'rb': 'ruby',
    'pl': 'perl'
}
COMMENT_REGEX_BY_LANG = {
    'javascript': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\'|`(?:(?:\\.|[^`\\])*)`)|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'typescript': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\'|`(?:(?:\\.|[^`\\])*)`)|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'go': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'php': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|//.*|#.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'csharp': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\'|`(?:(?:\\.|[^`\\])*)`)|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'kotlin': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'swift': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'rust': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'dart': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\'|`(?:(?:\\.|[^`\\])*)`)|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'scala': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'sql': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|--.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'r': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|#.*', re.MULTILINE),
    'lua': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|--\\[\\[[\\s\\S]*?\\]\\]|--.*', re.MULTILINE),
    'powershell': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|#.*|<#[\\s\\S]*?#>', re.MULTILINE),
    'yaml': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\')|#.*', re.MULTILINE),
    'java': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\'|`(?:(?:\\.|[^`\\])*)`)|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'c': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\'|`(?:(?:\\.|[^`\\])*)`)|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'cpp': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\'|`(?:(?:\\.|[^`\\])*)`)|//.*|/\*[\s\S]*?\*/', re.MULTILINE),
    'python': re.compile(r'#.*|\'\'\'[\s\S]*?\'\'\'|"""[\s\S]*?"""', re.MULTILINE),
    'html': re.compile(r'<!--[\s\S]*?-->', re.MULTILINE),
    'xml': re.compile(r'<!--[\s\S]*?-->', re.MULTILINE),
    'css': re.compile(r'/\*[\s\S]*?\*/', re.MULTILINE),
    'shellscript': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\'|`(?:(?:\\.|[^`\\])*)`)|#.*|=begin[\s\S]*?=end', re.MULTILINE),
    'ruby': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\'|`(?:(?:\\.|[^`\\])*)`)|#.*|=begin[\s\S]*?=end', re.MULTILINE),
    'perl': re.compile(r'("(?:(?:\\.|[^"\\])*)"|\'(?:(?:\\.|[^\'\\])*)\'|`(?:(?:\\.|[^`\\])*)`)|#.*|=begin[\s\S]*?=end', re.MULTILINE)
}

PYTHON_BLOCK_COMMENT = re.compile(r'\'\'\'[\s\S]*?\'\'\'|"""[\s\S]*?"""', re.MULTILINE)
PYTHON_LINE_COMMENT = re.compile(
    r'([rRuUfFbB]{0,2}"(?:(?:\\.|[^"\\])*)"|[rRuUfFbB]{0,2}\'(?:(?:\\.|[^\'\\])*)\')|#.*',
    re.MULTILINE
)


def del_slash(dirs):
    """
    删除目录字符串末尾的 “/”。

    说明：
        该函数仅处理以 “/” 结尾的路径（偏 POSIX 风格）。项目中同时提供了
        normalize_paths 用于跨平台的路径归一化；此函数保留用于兼容旧调用。

    Args:
        dirs: 目录列表

    Returns:
        处理后的目录列表
    """
    no_slash_dirs = []
    for dir_ in dirs:
        if dir_[-1] == '/':
            no_slash_dirs.append(dir_[: -1])
        else:
            no_slash_dirs.append(dir_)
    return no_slash_dirs


def normalize_paths(paths):
    """
    将路径转为绝对路径并去重。

    - 会移除末尾路径分隔符（os.sep）
    - 去重时按大小写不敏感（适配 Windows 路径）

    Args:
        paths: 路径列表（文件或目录）

    Returns:
        归一化后的路径列表（保持首次出现顺序）
    """
    if not paths:
        return []
    normalized = []
    for item in paths:
        if not item:
            continue
        normalized.append(abspath(item).rstrip(os.sep))
    unique = []
    seen = set()
    for item in normalized:
        key = os.path.normcase(item)
        if key in seen:
            continue
        seen.add(key)
        unique.append(item)
    return unique


def read_gitignore_excludes(indirs):
    """
    读取指定目录下的 .gitignore，并提取可解析为“静态路径”的排除项。

    说明：
        - 会忽略通配符/字符集等模式（如 *, ?, []）
        - 会忽略否定规则（以 ! 开头）
        - 仅保留在磁盘上真实存在的路径

    Args:
        indirs: 源码目录列表

    Returns:
        排除路径列表（绝对路径、去重）
    """
    excludes = []
    for indir in indirs:
        root = abspath(indir)
        if not os.path.isdir(root):
            continue
        for current_root, dirs, files in os.walk(root):
            gitignore_name = '.gitignore'
            if gitignore_name not in files:
                continue
            gitignore_path = os.path.join(current_root, gitignore_name)
            try:
                with codecs.open(gitignore_path, encoding='utf-8', errors='ignore') as fp:
                    lines = fp.readlines()
            except OSError:
                continue
            for raw in lines:
                line = raw.strip()
                if not line:
                    continue
                if line.startswith('#'):
                    continue
                if line.startswith('!'):
                    continue
                line = line.replace('\\', '/')
                if line.startswith('/'):
                    line = line[1:]
                if any(ch in line for ch in ('*', '?', '[')):
                    continue
                if line.endswith('/'):
                    line = line[:-1]
                if not line:
                    continue
                path = abspath(os.path.join(current_root, line))
                if os.path.exists(path):
                    excludes.append(path)
    return normalize_paths(excludes)


def is_binary_file(file_path):
    """
    粗略判断文件是否为二进制文件。

    Args:
        file_path: 文件路径

    Returns:
        bool：疑似二进制则为 True
    """
    try:
        with open(file_path, 'rb') as fp:
            chunk = fp.read(2048)
        return b'\x00' in chunk
    except OSError:
        return True


def decode_content(file_path, encoding):
    """
    读取源码文件内容并按编码解码为字符串。

    Args:
        file_path: 文件路径
        encoding: 指定编码；若为 'auto' 或空则尝试多种常见编码

    Returns:
        解码后的文本内容（失败时返回空字符串）
    """
    if encoding and encoding != 'auto':
        with codecs.open(file_path, encoding=encoding, errors='ignore') as fp:
            return fp.read()
    try:
        with open(file_path, 'rb') as fp:
            data = fp.read()
    except OSError:
        return ''
    for name in ('utf-8-sig', 'utf-8', 'gb18030', 'gbk'):
        try:
            return data.decode(name)
        except UnicodeDecodeError:
            continue
    return data.decode('utf-8', errors='ignore')


def get_language_by_extension(file_path):
    """
    根据文件后缀推断语言标识。

    Args:
        file_path: 文件路径

    Returns:
        语言字符串（未识别返回空字符串）
    """
    ext = os.path.splitext(file_path)[1].lower().lstrip('.')
    return LANGUAGE_BY_EXT.get(ext, '')


def strip_comments(content, language):
    """
    按语言规则移除注释内容。

    说明：
        - Python：优先移除三引号块，再移除行注释；并尽量保留字符串字面量中的 '#'
        - 其它语言：尽量保留字符串字面量中的注释符号

    Args:
        content: 源码文本
        language: 语言标识（由后缀推断）

    Returns:
        去除注释后的文本
    """
    if language == 'python':
        content = PYTHON_BLOCK_COMMENT.sub('', content)
        def replacer(match):
            group = match.group(1)
            return group if group else ''
        return PYTHON_LINE_COMMENT.sub(replacer, content)
    pattern = COMMENT_REGEX_BY_LANG.get(language)
    if not pattern:
        return content
    if language in ('html', 'xml', 'css'):
        return pattern.sub('', content)
    def replacer(match):
        group = match.group(1)
        return group if group else ''
    return pattern.sub(replacer, content)


def filter_lines(content, language, skip_blank_lines, skip_comment_lines, comment_chars):
    """
    将源码内容按规则过滤为“需要写入文档”的行列表。

    Args:
        content: 源码文本
        language: 语言标识；为空则按 comment_chars 做前缀判断
        skip_blank_lines: 是否过滤空行
        skip_comment_lines: 是否过滤注释行（或注释块）
        comment_chars: 注释前缀列表（language 为空时使用）

    Returns:
        过滤后的行列表
    """
    text = content
    if skip_comment_lines:
        if language:
            text = strip_comments(text, language)
        else:
            lines = []
            for line in text.splitlines():
                trimmed = line.lstrip()
                if any(trimmed.startswith(ch) for ch in comment_chars):
                    continue
                lines.append(line)
            text = '\n'.join(lines)
    lines = text.splitlines()
    if skip_blank_lines:
        lines = [line for line in lines if line.strip()]
    return lines


class CodeFinder(object):
    """
    递归扫描目录，收集指定后缀的代码文件。

    会跳过：
        - 隐藏文件/目录（以 . 开头）
        - 指定的目录名/文件名
        - excludes 命中的文件或目录
        - 疑似二进制文件
    """
    def __init__(self, exts=None, skip_dir_names=None, skip_file_names=None):
        """
        Args:
            exts: 后缀列表（如 ['py', 'js']），默认为 ['py']
            skip_dir_names: 需要跳过的目录名列表
            skip_file_names: 需要跳过的文件名列表
        """
        self.exts = exts if exts else ['py']
        self.skip_dir_names = skip_dir_names if skip_dir_names else []
        self.skip_file_names = skip_file_names if skip_file_names else []

    def is_code(self, file):
        lowered = file.lower()
        for ext in self.exts:
            ext = ext.lower().lstrip('.')
            if lowered.endswith('.' + ext):
                return True
        return False

    @staticmethod
    def is_hidden_file(file):
        """
        是否是隐藏文件（以 '.' 开头）。
        """
        return file[0] == '.'

    @staticmethod
    def should_be_excluded(file, excludes=None):
        """
        判断路径是否在排除列表中。

        Args:
            file: 绝对路径（文件或目录）
            excludes: 绝对路径列表；若包含目录，则其子路径也会被排除

        Returns:
            bool：需要排除则为 True
        """
        if not excludes:
            return False
        if not isinstance(excludes, list):
            excludes = [excludes]
        should_be_excluded = False
        for exclude in excludes:
            if file == exclude:
                should_be_excluded = True
                break
            prefix = exclude + os.sep
            if file.startswith(prefix):
                should_be_excluded = True
                break
        return should_be_excluded

    def find(self, indir, excludes=None):
        """
        查找目录下所有符合后缀的代码文件。

        Args:
            indir: 需要扫描的目录
            excludes: 排除文件或目录（绝对路径）

        Returns:
            代码文件列表（绝对路径）
        """
        files = []
        for entry in scandir(indir):
            entry_name = entry.name
            entry_path = abspath(entry.path)
            if self.is_hidden_file(entry_name):
                continue
            if entry.is_dir() and entry_name in self.skip_dir_names:
                continue
            if entry.is_file() and entry_name in self.skip_file_names:
                continue
            if self.should_be_excluded(entry_path, excludes):
                continue
            if entry.is_file():
                if is_binary_file(entry_path):
                    continue
                if self.is_code(entry_name):
                    files.append(entry_path)
                continue
            for file in self.find(entry_path, excludes=excludes):
                files.append(file)
        logger.debug('在%s目录下找到%d个代码文件.', indir, len(files))
        return files


class CodeWriter(object):
    """
    将源码文件按行写入 docx 文档。

    支持配置：
        - 页眉标题
        - 字体、字号、段前/段后/行距
        - 空行/注释过滤（与 filter_lines 一致）
        - docx 模板
        - 页眉字体独立设置
        - 页脚页码格式自定义
    """
    def __init__(
            self, font_name='宋体',
            font_size=10.5, space_before=0.0,
            space_after=2.3, line_spacing=10.5,
            command_chars=None, document=None,
            template_path=None, skip_blank_lines=True,
            skip_comment_lines=True, encoding='utf-8',
            header_font='宋体', header_font_size=10.5,
            page_number_format='第 {page} 页 共 {total} 页'
    ):
        Document, Pt, WD_PARAGRAPH_ALIGNMENT = load_docx_dependencies()
        self.font_name = font_name
        self.font_size = font_size
        self.space_before = space_before
        self.space_after = space_after
        self.line_spacing = line_spacing
        self.command_chars = command_chars if command_chars else DEFAULT_COMMENT_CHARS
        self.skip_blank_lines = skip_blank_lines
        self.skip_comment_lines = skip_comment_lines
        self.encoding = encoding
        self.header_font = header_font
        self.header_font_size = header_font_size
        self.page_number_format = page_number_format
        self._Pt = Pt
        self._WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT
        self.document = document if document else create_document(template_path, Document)
        self.total_lines = 0

    @staticmethod
    def is_blank_line(line):
        """
        判断是否是空行。
        """
        return not bool(line)

    def is_comment_line(self, line):
        line = line.lstrip()
        is_comment = False
        for comment_char in self.command_chars:
            if line.startswith(comment_char):
                is_comment = True
                break
        return is_comment

    def write_header(self, title):
        """
        写入页眉标题,使用独立的页眉字体设置。
        """
        from docx.oxml.ns import qn
        
        paragraph = self.document.sections[0].header.paragraphs[0]
        paragraph.alignment = self._WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(title)
        run.font.name = self.header_font
        run.font.size = self._Pt(self.header_font_size)
        # 为中文字体设置东亚字体属性
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.header_font)
        return self
    
    def write_footer(self):
        """
        添加页脚页码。
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        section = self.document.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = self._WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 创建简单的页码文本（动态页码需要通过域代码）
        # 如果格式包含 {page} 和 {total}，使用域代码
        if '{page}' in self.page_number_format and '{total}' in self.page_number_format:
            parts = self.page_number_format.split('{page}')
            if parts[0]:
                run = footer_para.add_run(parts[0])
                run.font.name = self.header_font
                run.font.size = self._Pt(self.header_font_size)
            
            # 添加当前页码域
            fld_simple = OxmlElement('w:fldSimple')
            fld_simple.set(qn('w:instr'), 'PAGE')
            run_elem = OxmlElement('w:r')
            # 为域代码设置字体属性
            r_pr = OxmlElement('w:rPr')
            r_fonts = OxmlElement('w:rFonts')
            r_fonts.set(qn('w:ascii'), self.header_font)
            r_fonts.set(qn('w:eastAsia'), self.header_font)
            r_fonts.set(qn('w:hAnsi'), self.header_font)
            r_pr.append(r_fonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(self.header_font_size * 2)))
            r_pr.append(sz)
            sz_cs = OxmlElement('w:szCs')
            sz_cs.set(qn('w:val'), str(int(self.header_font_size * 2)))
            r_pr.append(sz_cs)
            run_elem.append(r_pr)
            fld_simple.append(run_elem)
            footer_para._element.append(fld_simple)
            
            if len(parts) > 1 and '{total}' in parts[1]:
                mid_parts = parts[1].split('{total}')
                if mid_parts[0]:
                    run = footer_para.add_run(mid_parts[0])
                    run.font.name = self.header_font
                    run.font.size = self._Pt(self.header_font_size)
                
                # 添加总页数域
                fld_simple2 = OxmlElement('w:fldSimple')
                fld_simple2.set(qn('w:instr'), 'NUMPAGES')
                run_elem2 = OxmlElement('w:r')
                # 为域代码设置字体属性
                r_pr2 = OxmlElement('w:rPr')
                r_fonts2 = OxmlElement('w:rFonts')
                r_fonts2.set(qn('w:ascii'), self.header_font)
                r_fonts2.set(qn('w:eastAsia'), self.header_font)
                r_fonts2.set(qn('w:hAnsi'), self.header_font)
                r_pr2.append(r_fonts2)
                sz2 = OxmlElement('w:sz')
                sz2.set(qn('w:val'), str(int(self.header_font_size * 2)))
                r_pr2.append(sz2)
                sz_cs2 = OxmlElement('w:szCs')
                sz_cs2.set(qn('w:val'), str(int(self.header_font_size * 2)))
                r_pr2.append(sz_cs2)
                run_elem2.append(r_pr2)
                fld_simple2.append(run_elem2)
                footer_para._element.append(fld_simple2)
                
                if len(mid_parts) > 1:
                    run = footer_para.add_run(mid_parts[1])
                    run.font.name = self.header_font
                    run.font.size = self._Pt(self.header_font_size)
        else:
            # 简单文本格式
            run = footer_para.add_run(self.page_number_format)
            run.font.name = self.header_font
            run.font.size = self._Pt(self.header_font_size)
        
        return self

    def write_file(self, file):
        """
        将单个文件内容按行追加到文档中。
        """
        content = decode_content(file, self.encoding)
        language = get_language_by_extension(file)
        lines = filter_lines(
            content, language,
            self.skip_blank_lines,
            self.skip_comment_lines,
            self.command_chars
        )
        for line in lines:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.space_before = self._Pt(self.space_before)
            paragraph.paragraph_format.space_after = self._Pt(self.space_after)
            paragraph.paragraph_format.line_spacing = self._Pt(self.line_spacing)
            run = paragraph.add_run(line.rstrip())
            run.font.name = self.font_name
            run.font.size = self._Pt(self.font_size)
            self.total_lines += 1
        return self
    
    def get_total_lines(self):
        """
        获取已写入的总行数。
        """
        return self.total_lines

    def save(self, file):
        self.document.save(file)


def load_docx_dependencies():
    """
    延迟导入 python-docx 依赖，避免在仅扫描时强制安装。

    Returns:
        (Document, Pt, WD_PARAGRAPH_ALIGNMENT)

    Raises:
        RuntimeError: 未安装 python-docx 或安装了错误的 docx 包
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        return Document, Pt, WD_PARAGRAPH_ALIGNMENT
    except Exception as exc:
        raise RuntimeError('未检测到可用的python-docx，请安装python-docx并卸载docx包') from exc


def create_document(template_path, Document):
    """
    创建 docx 文档对象。

    Args:
        template_path: 模板文件路径；为空则创建空文档
        Document: python-docx 的 Document 类
    """
    if template_path:
        return Document(template_path)
    return Document()


def normalize_items(text):
    """
    将多行/逗号/分号分隔的文本整理为字符串列表。

    Args:
        text: 原始输入文本

    Returns:
        非空条目列表
    """
    items = []
    for line in text.splitlines():
        for part in line.replace(';', ',').split(','):
            item = part.strip()
            if item:
                items.append(item)
    return items


def normalize_exts(exts):
    """
    规范化后缀列表（去空白、去掉前导 '.'）。

    Args:
        exts: 后缀列表（可包含 '.py' 形式）

    Returns:
        规范化后的后缀列表
    """
    items = []
    for ext in exts:
        item = ext.strip()
        if not item:
            continue
        if item.startswith('.'):
            item = item[1:]
        items.append(item)
    return items


def collect_code_files(indirs, exts, excludes, skip_dir_names=None, skip_file_names=None):
    """
    收集所有代码文件路径。

    Args:
        indirs: 源码目录列表
        exts: 后缀列表
        excludes: 排除路径列表（绝对路径）
        skip_dir_names: 跳过目录名列表
        skip_file_names: 跳过文件名列表

    Returns:
        文件路径列表（绝对路径）
    """
    finder = CodeFinder(exts, skip_dir_names=skip_dir_names, skip_file_names=skip_file_names)
    files = []
    for indir in indirs:
        files.extend(finder.find(indir, excludes=excludes))
    return files


def collect_all_file_extensions(indirs, excludes, skip_dir_names=None, skip_file_names=None):
    """
    扫描目录并收集出现过的文件后缀（用于 GUI 的“自动识别后缀”）。

    Args:
        indirs: 源码目录列表
        excludes: 排除路径列表（绝对路径）
        skip_dir_names: 跳过目录名列表
        skip_file_names: 跳过文件名列表

    Returns:
        排序后的后缀列表
    """
    extensions = set()
    skip_dir_names = set(skip_dir_names or [])
    skip_file_names = set(skip_file_names or [])
    for indir in indirs:
        for root, dirs, files in os.walk(indir):
            root_path = abspath(root)
            if CodeFinder.should_be_excluded(root_path, excludes):
                dirs[:] = []
                continue
            dirs[:] = [d for d in dirs if not d.startswith('.') and d not in skip_dir_names]
            for name in files:
                if name.startswith('.'):
                    continue
                if name in skip_file_names:
                    continue
                file_path = abspath(os.path.join(root, name))
                if CodeFinder.should_be_excluded(file_path, excludes):
                    continue
                if is_binary_file(file_path):
                    continue
                ext = os.path.splitext(name)[1].lower().lstrip('.')
                if ext:
                    extensions.add(ext)
    return sorted(extensions)


def count_total_lines(files, skip_blank_lines, skip_comment_lines, comment_chars, encoding):
    """
    统计文件的总行数，并分别统计代码行数和注释行数。
    
    Args:
        files: 文件路径列表
        skip_blank_lines: 是否过滤空行
        skip_comment_lines: 是否过滤注释
        comment_chars: 注释前缀列表
        encoding: 文件编码
    
    Returns:
        dict: 包含以下键值：
            - total_lines: 总行数（过滤后）
            - code_lines: 代码行数（非注释、非空行）
            - comment_lines: 注释行数
            - blank_lines: 空行数
            - raw_lines: 原始总行数（未过滤）
    """
    total_filtered = 0
    code_lines = 0
    comment_lines = 0
    blank_lines = 0
    raw_lines = 0
    
    for file in files:
        content = decode_content(file, encoding)
        language = get_language_by_extension(file)
        
        # 统计原始行数
        raw_file_lines = content.splitlines()
        raw_lines += len(raw_file_lines)
        
        # 统计注释行（去除注释后的行数差）
        if language:
            stripped = strip_comments(content, language)
            stripped_lines = stripped.splitlines()
            
            # 统计空行
            for line in raw_file_lines:
                if not line.strip():
                    blank_lines += 1
            
            # 统计注释行（原始行数 - 去除注释后的行数 - 空行）
            # 这是一个近似值，因为注释可能和代码在同一行
            lines_after_comment_removal = len([l for l in stripped_lines if l.strip()])
            file_blank = sum(1 for line in raw_file_lines if not line.strip())
            file_comment = len(raw_file_lines) - lines_after_comment_removal - file_blank
            comment_lines += file_comment
            
            # 统计代码行
            code_lines += lines_after_comment_removal
        else:
            # 无法识别语言时，使用简单的前缀匹配
            for line in raw_file_lines:
                stripped = line.strip()
                if not stripped:
                    blank_lines += 1
                elif any(stripped.startswith(ch) for ch in comment_chars):
                    comment_lines += 1
                else:
                    code_lines += 1
        
        # 统计过滤后的行数
        lines = filter_lines(
            content, language,
            skip_blank_lines,
            skip_comment_lines,
            comment_chars
        )
        total_filtered += len(lines)
    
    return {
        'total_lines': total_filtered,
        'code_lines': code_lines,
        'comment_lines': comment_lines,
        'blank_lines': blank_lines,
        'raw_lines': raw_lines
    }


def generate_code_doc(
        title, indirs, exts, comment_chars,
        font_name, font_size, space_before,
        space_after, line_spacing, excludes,
        outfile, template_path=None,
        skip_blank_lines=True, skip_comment_lines=True,
        encoding='utf-8', skip_dir_names=None, skip_file_names=None,
        header_font='宋体', header_font_size=10.5,
        page_number_format='第 {page} 页 共 {total} 页',
        selected_files=None, pages_60_mode=False
):
    """
    生成 docx 源代码文档。

    Args:
        title: 页眉标题
        indirs: 源码目录列表
        exts: 后缀列表
        comment_chars: 注释前缀列表（language 未识别时使用）
        font_name/font_size/space_before/space_after/line_spacing: 排版参数
        excludes: 排除路径列表
        outfile: 输出 docx 路径
        template_path: 模板 docx 路径
        skip_blank_lines: 是否过滤空行
        skip_comment_lines: 是否过滤注释
        encoding: 源码文件编码；支持 'auto'
        skip_dir_names/skip_file_names: 跳过目录名/文件名列表
        header_font: 页眉字体
        header_font_size: 页眉字号
        page_number_format: 页码格式
        selected_files: 用户选择的文件列表（如果为 None 则使用所有文件）
        pages_60_mode: 是否启用60页模式（前30页+后30页）

    Returns:
        dict：包含 file_count、outfile、total_lines
    """
    if not indirs:
        indirs = DEFAULT_INDIRS
    if not exts:
        exts = DEFAULT_EXTS
    else:
        exts = normalize_exts(exts)
    if not comment_chars:
        comment_chars = DEFAULT_COMMENT_CHARS
    indirs = [abspath(indir) for indir in indirs]
    excludes = normalize_paths(excludes)
    if outfile:
        excludes = normalize_paths(excludes + [outfile])
    if not skip_dir_names:
        skip_dir_names = DEFAULT_SKIP_DIRS
    if not skip_file_names:
        skip_file_names = DEFAULT_SKIP_FILES
    
    files = collect_code_files(indirs, exts, excludes, skip_dir_names, skip_file_names)
    
    # 如果指定了 selected_files，则只使用选中的文件，并保持其顺序
    if selected_files:
        # 创建文件集合用于快速查找
        available_files_set = set(files)
        # 按照 selected_files 的顺序，只保留实际存在的文件
        files = [f for f in selected_files if f in available_files_set]
    
    # 统计总行数
    total_lines_info = count_total_lines(files, skip_blank_lines, skip_comment_lines, comment_chars, encoding)
    total_lines = total_lines_info['total_lines']
    
    writer = CodeWriter(
        command_chars=comment_chars,
        font_name=font_name,
        font_size=font_size,
        space_before=space_before,
        space_after=space_after,
        line_spacing=line_spacing,
        template_path=template_path,
        skip_blank_lines=skip_blank_lines,
        skip_comment_lines=skip_comment_lines,
        encoding=encoding,
        header_font=header_font,
        header_font_size=header_font_size,
        page_number_format=page_number_format
    )
    writer.write_header(title)
    writer.write_footer()
    
    # 60页模式：只写入前30页和后30页（假设每页50行）
    if pages_60_mode and total_lines >= 3000:
        lines_per_page = 50
        first_30_lines = lines_per_page * 30  # 1500行
        last_30_lines = lines_per_page * 30   # 1500行
        
        current_line = 0
        for file in files:
            content = decode_content(file, encoding)
            language = get_language_by_extension(file)
            lines = filter_lines(
                content, language,
                skip_blank_lines,
                skip_comment_lines,
                comment_chars
            )
            
            for line in lines:
                if current_line < first_30_lines or current_line >= (total_lines_info['total_lines'] - last_30_lines):
                    paragraph = writer.document.add_paragraph()
                    paragraph.paragraph_format.space_before = writer._Pt(space_before)
                    paragraph.paragraph_format.space_after = writer._Pt(space_after)
                    paragraph.paragraph_format.line_spacing = writer._Pt(line_spacing)
                    run = paragraph.add_run(line.rstrip())
                    run.font.name = font_name
                    run.font.size = writer._Pt(font_size)
                current_line += 1
    else:
        # 正常模式：写入所有文件
        for file in files:
            writer.write_file(file)
    
    writer.save(outfile)
    return {
        'file_count': len(files),
        'outfile': outfile,
        'total_lines': total_lines_info
    }
