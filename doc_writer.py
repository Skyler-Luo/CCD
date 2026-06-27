# -*- coding: utf-8 -*-
"""
文档写入模块
负责将代码写入docx文档，支持格式化和页眉页脚
"""
from code_processor import (
    DEFAULT_COMMENT_CHARS, decode_content,
    get_language_by_extension, filter_lines
)
from docx.oxml.ns import qn


def load_docx_dependencies():
    """
    延迟导入 python-docx 依赖，避免在仅扫描时强制安装
    
    Returns:
        (Document, Pt, WD_PARAGRAPH_ALIGNMENT)
        
    Raises:
        RuntimeError: 未安装 python-docx 或安装了错误的 docx 包
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        return Document, Pt, WD_PARAGRAPH_ALIGNMENT
    except Exception as exc:
        raise RuntimeError('未检测到可用的python-docx，请安装python-docx并卸载docx包') from exc


def create_document(template_path, Document):
    """
    创建 docx 文档对象
    
    Args:
        template_path: 模板文件路径；为空则创建空文档
        Document: python-docx 的 Document 类
    """
    if template_path:
        return Document(template_path)
    return Document()


class CodeWriter:
    """
    将源码文件按行写入 docx 文档
    
    支持配置：
        - 页眉标题
        - 字体、字号、段前/段后/行距
        - 空行/注释过滤
        - docx 模板
        - 页眉和页脚字体独立设置
        - 页脚页码格式自定义
    """
    def __init__(
            self, font_name='Consolas',
            font_size=10.5, space_before=0.0,
            space_after=2.3, line_spacing=10.5,
            command_chars=None, document=None,
            template_path=None, skip_blank_lines=True,
            skip_comment_lines=True, encoding='utf-8',
            header_font='宋体', header_font_size=10.5,
            top_margin=2.5, bottom_margin=2.5,
            left_margin=2.5, right_margin=2.5,
            header_distance=1.5, footer_distance=1.75,
            page_number_format='第 {page} 页 共 {total} 页'
    ):
        Document, Pt, WD_PARAGRAPH_ALIGNMENT = load_docx_dependencies()
        self.font_name = font_name
        self.font_size = font_size
        self.space_before = space_before
        self.space_after = space_after
        self.line_spacing = line_spacing
        self.command_chars = command_chars if command_chars else DEFAULT_COMMENT_CHARS
        self.skip_blank_lines = skip_blank_lines
        self.skip_comment_lines = skip_comment_lines
        self.encoding = encoding
        self.header_font = header_font
        self.header_font_size = header_font_size
        self.page_number_format = page_number_format
        self._Pt = Pt
        self._WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT
        self.document = document if document else create_document(template_path, Document)
        
        # 设置页面边距（可自定义，默认为软著标准）
        from docx.shared import Cm
        for section in self.document.sections:
            section.top_margin = Cm(top_margin)           # 上边距
            section.bottom_margin = Cm(bottom_margin)     # 下边距
            section.left_margin = Cm(left_margin)         # 左边距
            section.right_margin = Cm(right_margin)       # 右边距
            section.header_distance = Cm(header_distance) # 页眉距离页面顶端
            section.footer_distance = Cm(footer_distance) # 页脚距离页面底端
        
        self.total_lines = 0

    @staticmethod
    def is_blank_line(line):
        """判断是否是空行"""
        return not bool(line.strip())

    def is_comment_line(self, line):
        """判断是否是注释行"""
        line = line.lstrip()
        return any(line.startswith(ch) for ch in self.command_chars)

    def write_header(self, title):
        """写入页眉标题，使用独立的页眉和页脚字体设置"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        section = self.document.sections[0]
        header = section.header
        
        # 清空页眉中的所有段落
        while len(header.paragraphs) > 0:
            p = header.paragraphs[0]
            p._element.getparent().remove(p._element)
        
        # 重新添加一个干净的段落
        paragraph = header.add_paragraph()
        paragraph.alignment = self._WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 设置段落格式 - 紧凑无间距
        pf = paragraph.paragraph_format
        pf.space_before = self._Pt(0)
        pf.space_after = self._Pt(0)
        pf.line_spacing = 1.0  # 单倍行距
        pf.widow_control = False
        pf.keep_together = True
        
        # 在XML层面设置段落属性
        pPr = paragraph._element.get_or_add_pPr()
        
        # 设置段落间距为0
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')  # 240 twips = 单倍行距
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        
        # 添加 suppress line numbers（不显示行号）
        suppress = OxmlElement('w:suppressLineNumbers')
        pPr.append(suppress)
        
        # 添加文本
        run = paragraph.add_run(title)
        run.font.name = self.header_font
        run.font.size = self._Pt(self.header_font_size)
        # 为中文字体设置东亚字体属性
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.header_font)
        
        return self
    
    def write_footer(self):
        """添加页脚页码"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        section = self.document.sections[0]
        footer = section.footer
        
        # 清空页脚中的所有段落
        for para in footer.paragraphs:
            para._element.getparent().remove(para._element)
        
        # 重新添加一个干净的段落
        footer_para = footer.add_paragraph()
        footer_para.alignment = self._WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 禁止页脚换行 - 多重设置确保不换行
        footer_para.paragraph_format.widow_control = False
        footer_para.paragraph_format.keep_together = True
        footer_para.paragraph_format.keep_with_next = False
        footer_para.paragraph_format.space_before = self._Pt(0)
        footer_para.paragraph_format.space_after = self._Pt(0)
        footer_para.paragraph_format.line_spacing = 1.0  # 单倍行距，与页眉保持一致
        
        # 在XML层面禁止换行
        pPr = footer_para._element.get_or_add_pPr()
        # 添加 keepLines 元素（保持行在一起）
        keep_lines = OxmlElement('w:keepLines')
        pPr.append(keep_lines)
        
        # 如果格式包含 {page} 和 {total}，使用域代码
        if '{page}' in self.page_number_format and '{total}' in self.page_number_format:
            parts = self.page_number_format.split('{page}')
            if parts[0]:
                run = footer_para.add_run(parts[0])
                run.font.name = self.header_font
                run.font.size = self._Pt(self.header_font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.header_font)
            
            # 添加当前页码域
            fld_simple = OxmlElement('w:fldSimple')
            fld_simple.set(qn('w:instr'), 'PAGE')
            run_elem = OxmlElement('w:r')
            # 为域代码设置字体属性
            r_pr = OxmlElement('w:rPr')
            r_fonts = OxmlElement('w:rFonts')
            r_fonts.set(qn('w:ascii'), self.header_font)
            r_fonts.set(qn('w:eastAsia'), self.header_font)
            r_fonts.set(qn('w:hAnsi'), self.header_font)
            r_pr.append(r_fonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(self.header_font_size * 2)))
            r_pr.append(sz)
            sz_cs = OxmlElement('w:szCs')
            sz_cs.set(qn('w:val'), str(int(self.header_font_size * 2)))
            r_pr.append(sz_cs)
            run_elem.append(r_pr)
            fld_simple.append(run_elem)
            footer_para._element.append(fld_simple)
            
            if len(parts) > 1 and '{total}' in parts[1]:
                mid_parts = parts[1].split('{total}')
                if mid_parts[0]:
                    run = footer_para.add_run(mid_parts[0])
                    run.font.name = self.header_font
                    run.font.size = self._Pt(self.header_font_size)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.header_font)
                
                # 添加总页数域
                fld_simple2 = OxmlElement('w:fldSimple')
                fld_simple2.set(qn('w:instr'), 'NUMPAGES')
                run_elem2 = OxmlElement('w:r')
                # 为域代码设置字体属性
                r_pr2 = OxmlElement('w:rPr')
                r_fonts2 = OxmlElement('w:rFonts')
                r_fonts2.set(qn('w:ascii'), self.header_font)
                r_fonts2.set(qn('w:eastAsia'), self.header_font)
                r_fonts2.set(qn('w:hAnsi'), self.header_font)
                r_pr2.append(r_fonts2)
                sz2 = OxmlElement('w:sz')
                sz2.set(qn('w:val'), str(int(self.header_font_size * 2)))
                r_pr2.append(sz2)
                sz_cs2 = OxmlElement('w:szCs')
                sz_cs2.set(qn('w:val'), str(int(self.header_font_size * 2)))
                r_pr2.append(sz_cs2)
                run_elem2.append(r_pr2)
                fld_simple2.append(run_elem2)
                footer_para._element.append(fld_simple2)
                
                if len(mid_parts) > 1:
                    run = footer_para.add_run(mid_parts[1])
                    run.font.name = self.header_font
                    run.font.size = self._Pt(self.header_font_size)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.header_font)
        else:
            # 简单文本格式
            run = footer_para.add_run(self.page_number_format)
            run.font.name = self.header_font
            run.font.size = self._Pt(self.header_font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.header_font)
        
        return self

    def write_file(self, file):
        """将单个文件内容按行追加到文档中"""
        content = decode_content(file, self.encoding)
        language = get_language_by_extension(file)
        lines = filter_lines(
            content, language,
            self.skip_blank_lines,
            self.skip_comment_lines,
            self.command_chars
        )
        for line in lines:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.space_before = self._Pt(self.space_before)
            paragraph.paragraph_format.space_after = self._Pt(self.space_after)
            paragraph.paragraph_format.line_spacing = self._Pt(self.line_spacing)
            run = paragraph.add_run(line.rstrip())
            run.font.name = self.font_name
            run.font.size = self._Pt(self.font_size)
            # 同时设置东亚字体，避免中文字符回退到默认字体导致混排
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            self.total_lines += 1
        return self
    
    def get_total_lines(self):
        """获取已写入的总行数"""
        return self.total_lines

    def save(self, file):
        """保存文档"""
        self.document.save(file)
